import io
import logging
from typing import List, Optional
import requests
from bs4 import BeautifulSoup
import PyPDF2
from docx import Document as DocxDocument
import chardet

from src.core.models import Document, DocumentType
from langchain.text_splitter import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parser for extracting text from various document formats."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """Initialize document parser with chunking configuration."""
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )
    
    def parse_document(self, document: Document) -> Document:
        """Parse document content based on its type."""
        try:
            if document.type == DocumentType.PDF:
                content = self._parse_pdf(document)
            elif document.type == DocumentType.DOCX:
                content = self._parse_docx(document)
            elif document.type == DocumentType.URL:
                content = self._parse_url(document)
            elif document.type == DocumentType.TEXT:
                content = self._parse_text(document)
            else:
                logger.warning(f"Unsupported document type: {document.type}")
                content = ""
            
            document.content = content
            return document
            
        except Exception as e:
            logger.error(f"Error parsing document {document.name}: {e}")
            document.content = ""
            return document
    
    def _parse_pdf(self, document: Document) -> str:
        """Extract text from PDF document."""
        try:
            raw_content = document.metadata.get('raw_content')
            if not raw_content:
                logger.error(f"No raw content found for PDF: {document.name}")
                return ""
            
            pdf_file = io.BytesIO(raw_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_content.append(text)
            
            full_text = "\n".join(text_content)
            logger.info(f"Extracted {len(full_text)} characters from PDF: {document.name}")
            return full_text
            
        except Exception as e:
            logger.error(f"Error parsing PDF {document.name}: {e}")
            return ""
    
    def _parse_docx(self, document: Document) -> str:
        """Extract text from DOCX document."""
        try:
            raw_content = document.metadata.get('raw_content')
            if not raw_content:
                logger.error(f"No raw content found for DOCX: {document.name}")
                return ""
            
            docx_file = io.BytesIO(raw_content)
            doc = DocxDocument(docx_file)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n".join(text_content)
            logger.info(f"Extracted {len(full_text)} characters from DOCX: {document.name}")
            return full_text
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {document.name}: {e}")
            return ""
    
    def _parse_url(self, document: Document) -> str:
        """Extract text content from URL."""
        try:
            url = document.source
            if not url or url == "":
                url = document.metadata.get('web_view_link', '')
            
            if not url:
                logger.error(f"No URL found for document: {document.name}")
                return ""
            
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            # Detect encoding
            if response.encoding is None:
                detected = chardet.detect(response.content)
                response.encoding = detected['encoding']
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            logger.info(f"Extracted {len(text)} characters from URL: {url}")
            return text
            
        except requests.RequestException as e:
            logger.error(f"Error fetching URL {document.source}: {e}")
            return ""
        except Exception as e:
            logger.error(f"Error parsing URL content: {e}")
            return ""
    
    def _parse_text(self, document: Document) -> str:
        """Extract text from plain text document."""
        try:
            raw_content = document.metadata.get('raw_content')
            if not raw_content:
                logger.error(f"No raw content found for text document: {document.name}")
                return ""
            
            # Try to decode as text
            if isinstance(raw_content, bytes):
                # Detect encoding
                detected = chardet.detect(raw_content)
                encoding = detected['encoding'] or 'utf-8'
                try:
                    text = raw_content.decode(encoding)
                except UnicodeDecodeError:
                    text = raw_content.decode('utf-8', errors='ignore')
            else:
                text = str(raw_content)
            
            logger.info(f"Extracted {len(text)} characters from text document: {document.name}")
            return text
            
        except Exception as e:
            logger.error(f"Error parsing text document {document.name}: {e}")
            return ""
    
    def chunk_document(self, document: Document) -> List[str]:
        """Split document content into chunks for embedding."""
        if not document.content:
            return []
        
        try:
            chunks = self.text_splitter.split_text(document.content)
            logger.info(f"Split document {document.name} into {len(chunks)} chunks")
            return chunks
        except Exception as e:
            logger.error(f"Error chunking document {document.name}: {e}")
            return []
    
    def process_documents(self, documents: List[Document]) -> List[Document]:
        """Process a list of documents, parsing and chunking them."""
        processed_docs = []
        
        for doc in documents:
            # Parse the document
            parsed_doc = self.parse_document(doc)
            
            # Generate chunks
            chunks = self.chunk_document(parsed_doc)
            
            # Store chunk IDs
            parsed_doc.chunk_ids = [f"{parsed_doc.id}_chunk_{i}" for i in range(len(chunks))]
            
            # Store chunks in metadata for later use
            parsed_doc.metadata['chunks'] = chunks
            
            processed_docs.append(parsed_doc)
            
        logger.info(f"Processed {len(processed_docs)} documents")
        return processed_docs